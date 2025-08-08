from flask import Flask, render_template, request, send_file, redirect, url_for, flash, abort
from docx import Document
from docx.shared import RGBColor  # для установки чёрного цвета
import os, re, io
from datetime import datetime
from zipfile import ZipFile

TEMPLATE_DIR = "templates_docs"
GENERATED_DIR = "generated"
ALLOWED_SUFFIX = ".docx"

os.makedirs(GENERATED_DIR, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret")

# Ищем всё вида { ... }
PLACEHOLDER_RE = re.compile(r"\{[^{}]+\}")

def list_templates():
    return sorted([f for f in os.listdir(TEMPLATE_DIR) if f.lower().endswith(ALLOWED_SUFFIX)])

def extract_placeholders(docx_paths):
    """Собираем уникальные плейсхолдеры из тела, таблиц и колонтитулов."""
    found = set()
    for path in docx_paths:
        doc = Document(path)
        # основной текст
        for p in doc.paragraphs:
            full = "".join(run.text for run in p.runs) or p.text
            if full:
                found.update(PLACEHOLDER_RE.findall(full))
        # таблицы
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full = "".join(run.text for run in p.runs) or p.text
                        if full:
                            found.update(PLACEHOLDER_RE.findall(full))
        # колонтитулы
        for section in doc.sections:
            for p in section.header.paragraphs:
                full = "".join(run.text for run in p.runs) or p.text
                if full:
                    found.update(PLACEHOLDER_RE.findall(full))
            for p in section.footer.paragraphs:
                full = "".join(run.text for run in p.runs) or p.text
                if full:
                    found.update(PLACEHOLDER_RE.findall(full))
    return sorted(found)

def replace_placeholders_preserve_runs(paragraph, mapping):
    """
    Бережная замена: сохраняем исходные runs (а значит — шрифт, размер, жирность и т.п.).
    Меняем ТОЛЬКО текст плейсхолдера на value и делаем его чёрным.
    Если плейсхолдер пересекает несколько runs, значение вставляем в начальный run,
    конечный run обрезаем, промежуточные очищаем — верстка абзаца не «пересобирается».
    """
    if not paragraph.runs:
        return

    run_texts = [r.text or "" for r in paragraph.runs]
    full = "".join(run_texts)
    if not full:
        return

    matches = list(PLACEHOLDER_RE.finditer(full))
    if not matches:
        return

    # Префиксные суммы длин для маппинга глобального индекса к (run_idx, offset)
    lengths = [len(t) for t in run_texts]
    cumul = []
    s = 0
    for L in lengths:
        cumul.append(s)
        s += L

    def locate(pos: int):
        """Глобальный индекс -> (run_idx, offset_in_run)"""
        i = 0
        # ищем последний cumul[i] <= pos
        while i + 1 < len(cumul) and cumul[i + 1] <= pos:
            i += 1
        return i, pos - cumul[i]

    # Идём с конца, чтобы индексы не съезжали
    for m in reversed(matches):
        ph_text = m.group(0)
        if ph_text not in mapping:
            continue
        value = mapping[ph_text]

        start, end = m.start(), m.end()
        si, so = locate(start)
        ei, eo = locate(end - 1)  # последний символ плейсхолдера

        if si == ei:
            # Плейсхолдер целиком в одном run
            r = paragraph.runs[si]
            t = r.text or ""
            before = t[:so]
            after = t[eo + 1:]
            r.text = before + value + after
            try:
                r.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
        else:
            # Плейсхолдер пересекает несколько run'ов
            r_start = paragraph.runs[si]
            r_end   = paragraph.runs[ei]

            t_start = r_start.text or ""
            t_end   = r_end.text or ""

            before = t_start[:so]
            after  = t_end[eo + 1:]

            # Вставляем значение в начальный run, сохраняется его шрифт/размер
            r_start.text = before + value
            try:
                r_start.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass

            # Конечный run — оставляем хвост после плейсхолдера
            r_end.text = after

            # Промежуточные run'ы чистим
            for idx in range(si + 1, ei):
                paragraph.runs[idx].text = ""

def replace_in_doc(doc, mapping):
    # основной текст
    for p in doc.paragraphs:
        replace_placeholders_preserve_runs(p, mapping)
    # таблицы
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_preserve_runs(p, mapping)
    # колонтитулы
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_placeholders_preserve_runs(p, mapping)
        for p in section.footer.paragraphs:
            replace_placeholders_preserve_runs(p, mapping)

@app.route("/")
def index():
    templates = list_templates()
    selected = request.args.getlist("t")
    selected = [s for s in selected if s in templates]
    placeholders = extract_placeholders([os.path.join(TEMPLATE_DIR, f) for f in selected]) if selected else []
    return render_template("index.html", templates=templates, selected=selected, placeholders=placeholders)

@app.route("/placeholders")
def placeholders():
    templates = list_templates()
    selected = request.args.getlist("t")
    selected = [s for s in selected if s in templates]
    placeholders = extract_placeholders([os.path.join(TEMPLATE_DIR, f) for f in selected]) if selected else []
    return render_template("index.html", templates=templates, selected=selected, placeholders=placeholders)

@app.route("/generate", methods=["POST"])
def generate():
    templates = list_templates()
    selected = request.form.getlist("selected_templates")
    selected = [s for s in selected if s in templates]
    if not selected:
        flash("Выберите хотя бы один шаблон.", "warning")
        return redirect(url_for("index"))

    mapping = {k[3:]: v for k, v in request.form.items() if k.startswith("ph:")}
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_files = []

    for name in selected:
        path = os.path.join(TEMPLATE_DIR, name)
        doc = Document(path)
        replace_in_doc(doc, mapping)
        base = os.path.splitext(name)[0]
        out_name = f"{base}__{stamp}.docx"
        out_path = os.path.join(GENERATED_DIR, out_name)
        doc.save(out_path)
        out_files.append(out_path)

    mem = io.BytesIO()
    with ZipFile(mem, "w") as z:
        for fp in out_files:
            z.write(fp, arcname=os.path.basename(fp))
    mem.seek(0)
    return send_file(mem, as_attachment=True, download_name=f"contracts_{stamp}.zip")

@app.route("/downloads")
def downloads():
    files = sorted([f for f in os.listdir(GENERATED_DIR) if f.lower().endswith(".docx")], reverse=True)
    return render_template("downloads.html", files=files)

@app.route("/download/<name>")
def download_file(name):
    if not name.lower().endswith(".docx"):
        abort(404)
    path = os.path.join(GENERATED_DIR, name)
    if not os.path.exists(path):
        abort(404)
    return send_file(path, as_attachment=True, download_name=name)

@app.errorhandler(404)
def not_found(e):
    return render_template("error.html", code=404, message="Страница не найдена"), 404

@app.errorhandler(500)
def server_error(e):
    return render_template("error.html", code=500, message="Внутренняя ошибка сервера"), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
